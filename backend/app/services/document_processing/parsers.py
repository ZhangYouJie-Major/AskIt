"""
文件解析器模块

提供文件解析的抽象基类和工厂类，支持多种文档格式的解析
"""
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Type

from .types import ParsedDocument
from .exceptions import UnsupportedFileTypeError, FileParseError


class FileParser(ABC):
    """文件解析器抽象基类
    
    定义文件解析的统一接口，所有具体解析器必须继承此类
    """
    
    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        """解析文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            ParsedDocument: 包含文本内容和元数据的对象
            
        Raises:
            FileParseError: 文件解析失败
        """
        pass
    
    def _validate_file_exists(self, file_path: str) -> Path:
        """验证文件是否存在
        
        Args:
            file_path: 文件路径
            
        Returns:
            Path: 文件路径对象
            
        Raises:
            FileParseError: 文件不存在
        """
        path = Path(file_path)
        if not path.exists():
            raise FileParseError(
                file_path=file_path,
                reason="文件不存在"
            )
        if not path.is_file():
            raise FileParseError(
                file_path=file_path,
                reason="路径不是文件"
            )
        return path


class PDFParser(FileParser):
    """PDF 文件解析器
    
    使用 pypdf 提取 PDF 文本内容和页码信息
    """
    
    def parse(self, file_path: str) -> ParsedDocument:
        """解析 PDF 文件
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            ParsedDocument: 包含文本内容和页码信息的对象
            
        Raises:
            FileParseError: 文件解析失败
        """
        # 验证文件存在
        path = self._validate_file_exists(file_path)
        
        try:
            # 导入 pypdf 库
            from pypdf import PdfReader
            
            # 打开并读取 PDF 文件
            reader = PdfReader(str(path))
            
            # 获取页数
            page_count = len(reader.pages)
            
            # 提取所有页面的文本内容
            content_parts = []
            page_info = {}
            
            for page_num, page in enumerate(reader.pages, start=1):
                try:
                    # 提取页面文本
                    page_text = page.extract_text()
                    if page_text:
                        content_parts.append(page_text)
                        # 记录每页的字符数
                        page_info[f"page_{page_num}"] = {
                            "page_number": page_num,
                            "char_count": len(page_text)
                        }
                except Exception as e:
                    # 单页提取失败不应导致整个文档解析失败
                    # 记录警告并继续处理其他页面
                    page_info[f"page_{page_num}"] = {
                        "page_number": page_num,
                        "error": str(e)
                    }
            
            # 合并所有页面的文本内容
            content = "\n\n".join(content_parts)
            
            # 如果没有提取到任何文本，可能是扫描版 PDF 或损坏文件
            if not content.strip():
                raise FileParseError(
                    file_path=file_path,
                    reason="PDF 文件中未提取到文本内容（可能是扫描版或损坏文件）"
                )
            
            # 构建元数据
            metadata = {
                "page_count": page_count,
                "pages": page_info,
                "total_chars": len(content)
            }
            
            # 尝试提取 PDF 元数据（如标题、作者等）
            if reader.metadata:
                pdf_metadata = {}
                if reader.metadata.title:
                    pdf_metadata["title"] = reader.metadata.title
                if reader.metadata.author:
                    pdf_metadata["author"] = reader.metadata.author
                if reader.metadata.subject:
                    pdf_metadata["subject"] = reader.metadata.subject
                if reader.metadata.creator:
                    pdf_metadata["creator"] = reader.metadata.creator
                if pdf_metadata:
                    metadata["pdf_metadata"] = pdf_metadata
            
            return ParsedDocument(
                content=content,
                page_count=page_count,
                metadata=metadata
            )
            
        except FileParseError:
            # 重新抛出已经格式化的 FileParseError
            raise
        except Exception as e:
            # 捕获所有其他异常并包装为 FileParseError
            raise FileParseError(
                file_path=file_path,
                reason=f"PDF 解析失败: {type(e).__name__}",
                original_error=e
            )


class WordParser(FileParser):
    """Word 文档解析器
    
    使用 python-docx 提取 Word 文档文本内容和段落结构
    """
    
    def parse(self, file_path: str) -> ParsedDocument:
        """解析 Word 文档
        
        Args:
            file_path: Word 文档路径
            
        Returns:
            ParsedDocument: 包含文本内容和段落结构的对象
            
        Raises:
            FileParseError: 文件解析失败
        """
        # 验证文件存在
        path = self._validate_file_exists(file_path)
        
        try:
            # 导入 python-docx 库
            from docx import Document
            
            # 打开并读取 Word 文档
            doc = Document(str(path))
            
            # 提取所有段落的文本内容
            paragraphs = []
            paragraph_info = []
            
            for idx, para in enumerate(doc.paragraphs):
                # 获取段落文本
                para_text = para.text
                if para_text:  # 只保留非空段落
                    paragraphs.append(para_text)
                    # 记录段落信息
                    paragraph_info.append({
                        "index": idx,
                        "char_count": len(para_text),
                        "style": para.style.name if para.style else None
                    })
            
            # 合并所有段落的文本内容
            content = "\n\n".join(paragraphs)
            
            # 如果没有提取到任何文本，可能是空文档或损坏文件
            if not content.strip():
                raise FileParseError(
                    file_path=file_path,
                    reason="Word 文档中未提取到文本内容（可能是空文档或损坏文件）"
                )
            
            # 构建元数据
            metadata = {
                "paragraph_count": len(paragraphs),
                "paragraphs": paragraph_info,
                "total_chars": len(content)
            }
            
            # 尝试提取文档核心属性（如标题、作者等）
            try:
                core_props = doc.core_properties
                doc_metadata = {}
                if core_props.title:
                    doc_metadata["title"] = core_props.title
                if core_props.author:
                    doc_metadata["author"] = core_props.author
                if core_props.subject:
                    doc_metadata["subject"] = core_props.subject
                if core_props.created:
                    doc_metadata["created"] = str(core_props.created)
                if core_props.modified:
                    doc_metadata["modified"] = str(core_props.modified)
                if doc_metadata:
                    metadata["doc_properties"] = doc_metadata
            except Exception:
                # 如果无法提取属性，忽略错误继续
                pass
            
            return ParsedDocument(
                content=content,
                page_count=None,  # Word 文档没有固定页数概念
                metadata=metadata
            )
            
        except FileParseError:
            # 重新抛出已经格式化的 FileParseError
            raise
        except Exception as e:
            # 捕获所有其他异常并包装为 FileParseError
            raise FileParseError(
                file_path=file_path,
                reason=f"Word 文档解析失败: {type(e).__name__}",
                original_error=e
            )


class TextParser(FileParser):
    """纯文本文件解析器
    
    读取纯文本文件的完整内容
    """
    
    def parse(self, file_path: str) -> ParsedDocument:
        """解析纯文本文件
        
        Args:
            file_path: 文本文件路径
            
        Returns:
            ParsedDocument: 包含文本内容的对象
            
        Raises:
            FileParseError: 文件解析失败
        """
        # 验证文件存在
        path = self._validate_file_exists(file_path)
        
        try:
            # 尝试多种编码读取文件
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as f:
                        content = f.read()
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，抛出错误
            if content is None:
                raise FileParseError(
                    file_path=file_path,
                    reason=f"无法使用任何支持的编码读取文件（尝试过: {', '.join(encodings)}）"
                )
            
            # 如果文件为空，抛出错误
            if not content.strip():
                raise FileParseError(
                    file_path=file_path,
                    reason="文本文件为空"
                )
            
            # 统计行数
            line_count = content.count('\n') + 1
            
            # 构建元数据
            metadata = {
                "encoding": used_encoding,
                "line_count": line_count,
                "total_chars": len(content)
            }
            
            return ParsedDocument(
                content=content,
                page_count=None,  # 纯文本文件没有页数概念
                metadata=metadata
            )
            
        except FileParseError:
            # 重新抛出已经格式化的 FileParseError
            raise
        except Exception as e:
            # 捕获所有其他异常并包装为 FileParseError
            raise FileParseError(
                file_path=file_path,
                reason=f"文本文件解析失败: {type(e).__name__}",
                original_error=e
            )


class MarkdownParser(FileParser):
    """Markdown 文件解析器
    
    读取 Markdown 文件的完整内容，保留格式标记
    """
    
    def parse(self, file_path: str) -> ParsedDocument:
        """解析 Markdown 文件
        
        Args:
            file_path: Markdown 文件路径
            
        Returns:
            ParsedDocument: 包含文本内容和格式标记的对象
            
        Raises:
            FileParseError: 文件解析失败
        """
        # 验证文件存在
        path = self._validate_file_exists(file_path)
        
        try:
            # 尝试多种编码读取文件
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as f:
                        content = f.read()
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，抛出错误
            if content is None:
                raise FileParseError(
                    file_path=file_path,
                    reason=f"无法使用任何支持的编码读取文件（尝试过: {', '.join(encodings)}）"
                )
            
            # 如果文件为空，抛出错误
            if not content.strip():
                raise FileParseError(
                    file_path=file_path,
                    reason="Markdown 文件为空"
                )
            
            # 统计行数
            line_count = content.count('\n') + 1
            
            # 分析 Markdown 结构（标题、代码块等）
            lines = content.split('\n')
            heading_count = 0
            code_block_count = 0
            in_code_block = False
            
            for line in lines:
                # 统计标题
                if line.strip().startswith('#'):
                    heading_count += 1
                # 统计代码块
                if line.strip().startswith('```'):
                    if not in_code_block:
                        code_block_count += 1
                    in_code_block = not in_code_block
            
            # 构建元数据
            metadata = {
                "encoding": used_encoding,
                "line_count": line_count,
                "total_chars": len(content),
                "heading_count": heading_count,
                "code_block_count": code_block_count,
                "format": "markdown"
            }
            
            return ParsedDocument(
                content=content,
                page_count=None,  # Markdown 文件没有页数概念
                metadata=metadata
            )
            
        except FileParseError:
            # 重新抛出已经格式化的 FileParseError
            raise
        except Exception as e:
            # 捕获所有其他异常并包装为 FileParseError
            raise FileParseError(
                file_path=file_path,
                reason=f"Markdown 文件解析失败: {type(e).__name__}",
                original_error=e
            )


class FileParserFactory:
    """文件解析器工厂类
    
    根据文件扩展名返回对应的解析器实例
    """
    
    # 支持的文件类型到解析器类的映射
    _PARSERS: Dict[str, Type[FileParser]] = {
        "pdf": PDFParser,
        "docx": WordParser,
        "txt": TextParser,
        "md": MarkdownParser,
    }
    
    @classmethod
    def get_parser(cls, file_type: str) -> FileParser:
        """根据文件类型返回对应的解析器
        
        Args:
            file_type: 文件扩展名（不区分大小写）
            
        Returns:
            FileParser: 对应的解析器实例
            
        Raises:
            UnsupportedFileTypeError: 文件类型不支持
        """
        # 转换为小写以支持不区分大小写
        file_type_lower = file_type.lower()
        
        # 查找对应的解析器类
        parser_class = cls._PARSERS.get(file_type_lower)
        
        if parser_class is None:
            # 抛出不支持的文件类型错误
            raise UnsupportedFileTypeError(
                file_type=file_type,
                supported_types=list(cls._PARSERS.keys())
            )
        
        # 返回解析器实例
        return parser_class()
    
    @classmethod
    def get_supported_types(cls) -> list:
        """获取支持的文件类型列表
        
        Returns:
            list: 支持的文件扩展名列表
        """
        return list(cls._PARSERS.keys())
