"""
测试 WordParser 使用真实的 Word 文档
"""
import sys
from pathlib import Path
import tempfile

# 添加 backend 目录到路径
backend_dir = Path(__file__).parent
sys.path.insert(0, str(backend_dir))

# 直接导入类型和异常
import importlib.util

# 加载 types 模块
types_path = backend_dir / "app" / "services" / "document_processing" / "types.py"
spec = importlib.util.spec_from_file_location("types_module", types_path)
types_module = importlib.util.module_from_spec(spec)
spec.loader.exec_module(types_module)

# 加载 exceptions 模块
exceptions_path = backend_dir / "app" / "services" / "document_processing" / "exceptions.py"
spec = importlib.util.spec_from_file_location("exceptions_module", exceptions_path)
exceptions_module = importlib.util.module_from_spec(spec)
spec.loader.exec_module(exceptions_module)

# 加载 parsers 模块
parsers_path = backend_dir / "app" / "services" / "document_processing" / "parsers.py"
spec = importlib.util.spec_from_file_location("parsers_module", parsers_path)
parsers_module = importlib.util.module_from_spec(spec)

# 注入依赖
sys.modules['types_module'] = types_module
sys.modules['exceptions_module'] = exceptions_module

# 修改 parsers 模块的导入
parsers_code = parsers_path.read_text(encoding='utf-8')
parsers_code = parsers_code.replace(
    "from .types import ParsedDocument",
    "from types_module import ParsedDocument"
).replace(
    "from .exceptions import UnsupportedFileTypeError, FileParseError",
    "from exceptions_module import UnsupportedFileTypeError, FileParseError"
)

# 执行修改后的代码
exec(parsers_code, parsers_module.__dict__)

# 提取需要的类
WordParser = parsers_module.WordParser
ParsedDocument = types_module.ParsedDocument
FileParseError = exceptions_module.FileParseError


def create_test_word_document():
    """创建一个测试用的 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt
        
        # 创建文档
        doc = Document()
        
        # 添加标题
        doc.add_heading('测试文档标题', 0)
        
        # 添加段落
        doc.add_paragraph('这是第一段内容。包含一些测试文字。')
        doc.add_paragraph('这是第二段内容。包含更多的测试文字。')
        
        # 添加二级标题
        doc.add_heading('第一节', level=1)
        doc.add_paragraph('第一节的内容。这里有一些详细的描述。')
        
        # 添加三级标题
        doc.add_heading('子节 1.1', level=2)
        doc.add_paragraph('子节的内容。包含一些具体的信息。')
        
        # 添加列表
        doc.add_paragraph('项目 1', style='List Bullet')
        doc.add_paragraph('项目 2', style='List Bullet')
        doc.add_paragraph('项目 3', style='List Bullet')
        
        # 添加另一节
        doc.add_heading('第二节', level=1)
        doc.add_paragraph('第二节的内容。这是最后一段文字。')
        
        # 保存到临时文件
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc.save(temp_file.name)
        temp_file.close()
        
        return temp_file.name
    except ImportError:
        print("警告: python-docx 未安装，跳过 Word 文档创建")
        return None


def test_word_parser_with_real_document():
    """测试 WordParser 使用真实的 Word 文档"""
    print("测试: WordParser 使用真实的 Word 文档")
    
    # 创建测试文档
    doc_path = create_test_word_document()
    
    if doc_path is None:
        print("  ⚠ 跳过测试（python-docx 未安装）")
        return
    
    try:
        parser = WordParser()
        result = parser.parse(doc_path)
        
        # 验证返回类型
        assert isinstance(result, ParsedDocument), "返回值应该是 ParsedDocument"
        
        # 验证内容
        assert "测试文档标题" in result.content, "应包含标题"
        assert "第一段内容" in result.content, "应包含第一段"
        assert "第二段内容" in result.content, "应包含第二段"
        assert "第一节" in result.content, "应包含第一节标题"
        assert "第二节" in result.content, "应包含第二节标题"
        assert "项目 1" in result.content, "应包含列表项"
        
        # 验证元数据
        assert result.metadata is not None, "元数据不应为空"
        assert "paragraph_count" in result.metadata, "元数据应包含 paragraph_count"
        assert "paragraphs" in result.metadata, "元数据应包含 paragraphs"
        assert "total_chars" in result.metadata, "元数据应包含 total_chars"
        
        # 验证段落数量
        assert result.metadata["paragraph_count"] > 0, "段落数应大于 0"
        
        # 验证 page_count 为 None（Word 文档没有固定页数）
        assert result.page_count is None, "Word 文档的 page_count 应为 None"
        
        print("  ✓ WordParser 基本功能正常")
        print(f"  ✓ 段落数: {result.metadata['paragraph_count']}")
        print(f"  ✓ 字符数: {result.metadata['total_chars']}")
        print(f"  ✓ 内容预览: {result.content[:100]}...")
        
        # 如果有文档属性，打印出来
        if "doc_properties" in result.metadata:
            print(f"  ✓ 文档属性: {result.metadata['doc_properties']}")
        
        print("测试通过！\n")
        
    finally:
        # 清理临时文件
        import os
        if doc_path and os.path.exists(doc_path):
            os.unlink(doc_path)


def test_word_parser_empty_document():
    """测试 WordParser 处理空文档"""
    print("测试: WordParser 处理空文档")
    
    try:
        from docx import Document
        
        # 创建空文档
        doc = Document()
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc.save(temp_file.name)
        temp_file.close()
        
        try:
            parser = WordParser()
            try:
                result = parser.parse(temp_file.name)
                assert False, "空文档应该抛出 FileParseError"
            except FileParseError as e:
                assert "空" in str(e) or "未提取到" in str(e), "错误消息应提到文档为空或未提取到内容"
                print("  ✓ 空文档正确抛出 FileParseError")
        finally:
            import os
            os.unlink(temp_file.name)
        
        print("测试通过！\n")
        
    except ImportError:
        print("  ⚠ 跳过测试（python-docx 未安装）\n")


if __name__ == "__main__":
    print("=" * 60)
    print("WordParser 真实文档测试")
    print("=" * 60 + "\n")
    
    test_word_parser_with_real_document()
    test_word_parser_empty_document()
    
    print("=" * 60)
    print("所有测试完成！✓")
    print("=" * 60)
