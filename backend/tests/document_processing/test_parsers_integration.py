"""
集成测试：测试所有解析器通过工厂类的完整流程
"""
import sys
from pathlib import Path
import tempfile
import os

# 添加 backend 目录到路径
backend_dir = Path(__file__).parent
sys.path.insert(0, str(backend_dir))

# 直接导入类型和异常
import importlib.util

# 加载 types 模块
types_path = backend_dir / "app" / "services" / "document_processing" / "types.py"
spec = importlib.util.spec_from_file_location("types_module", types_path)
types_module = importlib.util.module_from_spec(spec)
spec.loader.exec_module(types_module)

# 加载 exceptions 模块
exceptions_path = backend_dir / "app" / "services" / "document_processing" / "exceptions.py"
spec = importlib.util.spec_from_file_location("exceptions_module", exceptions_path)
exceptions_module = importlib.util.module_from_spec(spec)
spec.loader.exec_module(exceptions_module)

# 加载 parsers 模块
parsers_path = backend_dir / "app" / "services" / "document_processing" / "parsers.py"
spec = importlib.util.spec_from_file_location("parsers_module", parsers_path)
parsers_module = importlib.util.module_from_spec(spec)

# 注入依赖
sys.modules['types_module'] = types_module
sys.modules['exceptions_module'] = exceptions_module

# 修改 parsers 模块的导入
parsers_code = parsers_path.read_text(encoding='utf-8')
parsers_code = parsers_code.replace(
    "from .types import ParsedDocument",
    "from types_module import ParsedDocument"
).replace(
    "from .exceptions import UnsupportedFileTypeError, FileParseError",
    "from exceptions_module import UnsupportedFileTypeError, FileParseError"
)

# 执行修改后的代码
exec(parsers_code, parsers_module.__dict__)

# 提取需要的类
FileParserFactory = parsers_module.FileParserFactory
ParsedDocument = types_module.ParsedDocument


def test_complete_workflow():
    """测试完整的文档解析工作流"""
    print("=" * 60)
    print("完整工作流测试")
    print("=" * 60 + "\n")
    
    test_files = []
    
    try:
        # 1. 创建测试文本文件
        print("1. 测试纯文本文件解析")
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("这是一个测试文本文件。\n包含多行内容。\n用于测试解析器。")
            txt_path = f.name
            test_files.append(txt_path)
        
        parser = FileParserFactory.get_parser("txt")
        result = parser.parse(txt_path)
        assert isinstance(result, ParsedDocument)
        assert "测试文本文件" in result.content
        assert result.metadata["encoding"] == "utf-8"
        print(f"  ✓ 文本文件解析成功")
        print(f"    - 字符数: {result.metadata['total_chars']}")
        print(f"    - 行数: {result.metadata['line_count']}\n")
        
        # 2. 创建测试 Markdown 文件
        print("2. 测试 Markdown 文件解析")
        markdown_content = """# 测试标题

## 第一节

这是**粗体**文字和*斜体*文字。

### 子节

- 列表项 1
- 列表项 2

```python
def test():
    print("Hello")
```

## 第二节

更多内容。
"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write(markdown_content)
            md_path = f.name
            test_files.append(md_path)
        
        parser = FileParserFactory.get_parser("md")
        result = parser.parse(md_path)
        assert isinstance(result, ParsedDocument)
        assert "# 测试标题" in result.content
        assert "**粗体**" in result.content
        assert "```python" in result.content
        assert result.metadata["format"] == "markdown"
        print(f"  ✓ Markdown 文件解析成功")
        print(f"    - 字符数: {result.metadata['total_chars']}")
        print(f"    - 标题数: {result.metadata['heading_count']}")
        print(f"    - 代码块数: {result.metadata['code_block_count']}\n")
        
        # 3. 创建测试 Word 文档
        print("3. 测试 Word 文档解析")
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading('测试 Word 文档', 0)
            doc.add_paragraph('这是第一段。')
            doc.add_paragraph('这是第二段。')
            doc.add_heading('第一节', level=1)
            doc.add_paragraph('第一节的内容。')
            
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
                docx_path = f.name
                test_files.append(docx_path)
            doc.save(docx_path)
            
            parser = FileParserFactory.get_parser("docx")
            result = parser.parse(docx_path)
            assert isinstance(result, ParsedDocument)
            assert "测试 Word 文档" in result.content
            assert "第一段" in result.content
            assert result.metadata["paragraph_count"] > 0
            print(f"  ✓ Word 文档解析成功")
            print(f"    - 字符数: {result.metadata['total_chars']}")
            print(f"    - 段落数: {result.metadata['paragraph_count']}\n")
            
        except ImportError:
            print("  ⚠ 跳过 Word 文档测试（python-docx 未安装）\n")
        
        # 4. 测试所有解析器返回的对象都符合规范
        print("4. 验证所有解析结果符合规范")
        for file_path in test_files:
            ext = Path(file_path).suffix[1:]  # 去掉点号
            if ext == "docx":
                try:
                    from docx import Document
                except ImportError:
                    continue
            
            parser = FileParserFactory.get_parser(ext)
            result = parser.parse(file_path)
            
            # 验证必填字段
            assert result.content, "content 不能为空"
            assert isinstance(result.metadata, dict), "metadata 必须是字典"
            assert "total_chars" in result.metadata, "metadata 必须包含 total_chars"
            
            print(f"  ✓ {ext.upper()} 解析结果符合规范")
        
        print("\n" + "=" * 60)
        print("所有集成测试通过！✓")
        print("=" * 60)
        
    finally:
        # 清理所有临时文件
        for file_path in test_files:
            if os.path.exists(file_path):
                os.unlink(file_path)


def test_requirements_validation():
    """验证需求 1.2, 1.3, 1.4, 1.6 的实现"""
    print("\n" + "=" * 60)
    print("需求验证测试")
    print("=" * 60 + "\n")
    
    test_files = []
    
    try:
        # 需求 1.2: Word 文档提取文本内容和段落结构
        print("需求 1.2: Word 文档提取文本内容和段落结构")
        try:
            from docx import Document
            
            doc = Document()
            doc.add_paragraph('段落 1')
            doc.add_paragraph('段落 2')
            doc.add_paragraph('段落 3')
            
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
                docx_path = f.name
                test_files.append(docx_path)
            doc.save(docx_path)
            
            parser = FileParserFactory.get_parser("docx")
            result = parser.parse(docx_path)
            
            # 验证提取了文本内容
            assert "段落 1" in result.content
            assert "段落 2" in result.content
            assert "段落 3" in result.content
            
            # 验证提取了段落结构
            assert "paragraph_count" in result.metadata
            assert "paragraphs" in result.metadata
            assert result.metadata["paragraph_count"] == 3
            
            print("  ✓ 需求 1.2 验证通过\n")
            
        except ImportError:
            print("  ⚠ 跳过需求 1.2（python-docx 未安装）\n")
        
        # 需求 1.3: 纯文本文件读取完整文本内容
        print("需求 1.3: 纯文本文件读取完整文本内容")
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("完整的文本内容\n第二行\n第三行")
            txt_path = f.name
            test_files.append(txt_path)
        
        parser = FileParserFactory.get_parser("txt")
        result = parser.parse(txt_path)
        
        # 验证读取了完整内容
        assert "完整的文本内容" in result.content
        assert "第二行" in result.content
        assert "第三行" in result.content
        
        print("  ✓ 需求 1.3 验证通过\n")
        
        # 需求 1.4: Markdown 文件读取完整文本内容并保留格式标记
        print("需求 1.4: Markdown 文件读取完整文本内容并保留格式标记")
        markdown_content = """# 标题

**粗体** 和 *斜体*

```code
block
```
"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write(markdown_content)
            md_path = f.name
            test_files.append(md_path)
        
        parser = FileParserFactory.get_parser("md")
        result = parser.parse(md_path)
        
        # 验证读取了完整内容
        assert "# 标题" in result.content
        
        # 验证保留了格式标记
        assert "**粗体**" in result.content
        assert "*斜体*" in result.content
        assert "```code" in result.content
        
        print("  ✓ 需求 1.4 验证通过\n")
        
        # 需求 1.6: 文件损坏或无法读取时返回描述性错误信息
        print("需求 1.6: 文件损坏或无法读取时返回描述性错误信息")
        
        # 测试不存在的文件
        parser = FileParserFactory.get_parser("txt")
        try:
            result = parser.parse("/nonexistent/file.txt")
            assert False, "应该抛出 FileParseError"
        except Exception as e:
            assert "不存在" in str(e) or "FileParseError" in str(type(e).__name__)
            print(f"  ✓ 不存在的文件返回描述性错误: {str(e)[:50]}...")
        
        # 测试空文件
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("")
            empty_path = f.name
            test_files.append(empty_path)
        
        parser = FileParserFactory.get_parser("txt")
        try:
            result = parser.parse(empty_path)
            assert False, "应该抛出 FileParseError"
        except Exception as e:
            assert "空" in str(e) or "FileParseError" in str(type(e).__name__)
            print(f"  ✓ 空文件返回描述性错误: {str(e)[:50]}...")
        
        print("  ✓ 需求 1.6 验证通过\n")
        
        print("=" * 60)
        print("所有需求验证通过！✓")
        print("=" * 60)
        
    finally:
        # 清理所有临时文件
        for file_path in test_files:
            if os.path.exists(file_path):
                os.unlink(file_path)


if __name__ == "__main__":
    test_complete_workflow()
    test_requirements_validation()
